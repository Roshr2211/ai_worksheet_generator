import streamlit as st
import json
import logging
import re
import uuid
import sqlite3
import pandas as pd
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from openai import OpenAI

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Database Class
class WorksheetDatabase:
    def __init__(self, db_path='worksheets.db'):
        """Initialize database connection and create table if not exists"""
        self.db_path = db_path
        self._create_table()

    def _create_table(self):
        """Create worksheets table if not exists"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS worksheets (
                    id TEXT PRIMARY KEY,
                    subject TEXT,
                    topic TEXT,
                    worksheet_content TEXT,
                    created_at DATETIME,
                    user_id TEXT
                )
            ''')
            conn.commit()

    def save_worksheet(self, subject, topic, worksheet_content, user_id=None):
        """Save a worksheet to the database"""
        worksheet_id = str(uuid.uuid4())
        created_at = datetime.now()

        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO worksheets 
                (id, subject, topic, worksheet_content, created_at, user_id) 
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                worksheet_id, 
                subject, 
                topic, 
                json.dumps(worksheet_content),  # Store as JSON string
                created_at, 
                user_id
            ))
            conn.commit()
        
        return worksheet_id

    def get_worksheet_by_id(self, worksheet_id):
        """Retrieve a specific worksheet by its ID"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM worksheets WHERE id = ?
            ''', (worksheet_id,))
            result = cursor.fetchone()
        
        if result:
            # Unpack the result
            id, subject, topic, worksheet_content, created_at, user_id = result
            return {
                'id': id,
                'subject': subject,
                'topic': topic,
                'worksheet_content': json.loads(worksheet_content),
                'created_at': created_at,
                'user_id': user_id
            }
        return None

    def list_worksheets(self, subject=None, limit=10):
        """List worksheets, optionally filtered by subject"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            if subject:
                cursor.execute('''
                    SELECT id, subject, topic, created_at 
                    FROM worksheets 
                    WHERE subject = ? 
                    ORDER BY created_at DESC 
                    LIMIT ?
                ''', (subject, limit))
            else:
                cursor.execute('''
                    SELECT id, subject, topic, created_at 
                    FROM worksheets 
                    ORDER BY created_at DESC 
                    LIMIT ?
                ''', (limit,))
            
            return cursor.fetchall()

# Initialize database
worksheet_db = WorksheetDatabase()

# Clean JSON Content
def clean_json_content(raw_content):
    """
    Comprehensively clean and prepare JSON content for parsing.
    Extremely robust method to extract a valid JSON object.
    """
    import re
    import json

    # Convert to string
    raw_content = str(raw_content)

    # Try multiple strategies to extract JSON
    def extract_json_strategies():
        strategies = [
            # Strategy 1: Regex to find JSON object with worksheet key
            lambda x: re.search(r'\{[^}]*"worksheet"\s*:\s*\[[^\]]+\][^}]*\}', x),
            
            # Strategy 2: Find between first { and last }
            lambda x: re.search(r'\{.*\}', x, re.DOTALL),
            
            # Strategy 3: Extract everything between first { and last }
            lambda x: x[x.find('{'):x.rfind('}')+1],
        ]

        for strategy in strategies:
            try:
                match = strategy(raw_content)
                if match:
                    # If it's a regex match, get the matched group
                    cleaned = match.group(0) if hasattr(match, 'group') else match
                    
                    # Additional cleaning
                    cleaned = re.sub(r',\s*}', '}', cleaned)  # Remove trailing commas
                    cleaned = re.sub(r'{\s*"', '{"', cleaned)  # Ensure proper key start
                    cleaned = cleaned.strip()

                    # Validate JSON structure
                    parsed = json.loads(cleaned)
                    
                    # Ensure worksheet key exists and has 9 elements
                    if (
                        "worksheet" in parsed and 
                        isinstance(parsed["worksheet"], list) and 
                        len(parsed["worksheet"]) == 9 and
                        all(isinstance(item, str) and item.strip() for item in parsed["worksheet"])
                    ):
                        return cleaned
            except (json.JSONDecodeError, Exception) as e:
                # Log or print the specific error if needed
                logger.warning(f"JSON extraction strategy failed: {e}")
        
        return None

    # Try to extract JSON
    cleaned_json = extract_json_strategies()
    
    if cleaned_json:
        return cleaned_json
    
    # Last resort: raise an error
    raise ValueError("Unable to extract valid JSON object")

def response(prompt_text, max_retries=3):
    progress_bar = st.progress(0)
    status_text = st.empty()

    for attempt in range(max_retries):
        try:
            status_text.info(f"Generating worksheet (Attempt {attempt + 1}/{max_retries})...")
            progress_bar.progress(min((attempt + 1) * 25, 75))
            
            logger.info(f"Attempt {attempt + 1}: Sending prompt")
            
            response = client.chat.completions.create(
                model="deepseek/deepseek-r1-zero:free",
                response_format={"type": "json_object"},
                messages=[
                    {
                        "role": "system", 
                        "content": "VERY IMPORTANT: Respond ONLY with a SINGLE, VALID JSON object. The object MUST have a 'worksheet' key with EXACTLY 9 string elements. NO additional text before or after the JSON. Ensure each element is a non-empty string."
                    },
                    {"role": "user", "content": prompt_text},
                ]
            )
            
            raw_content = response.choices[0].message.content
            logger.info(f"Raw response content: {raw_content}")
            
            status_text.info("Cleaning and parsing response...")
            progress_bar.progress(min((attempt + 1) * 40, 90))
            
            # Comprehensive cleaning
            cleaned_content = clean_json_content(raw_content)
            logger.info(f"Cleaned content: {cleaned_content}")
            
            # Parse cleaned content
            worksheet_data = json.loads(cleaned_content)
            
            progress_bar.progress(100)
            status_text.success("Worksheet generated successfully!")
            
            # Delay and clear progress indicators
            import time
            time.sleep(1)
            progress_bar.empty()
            status_text.empty()
            
            return worksheet_data["worksheet"]
            
        except (json.JSONDecodeError, ValueError) as e:
            logger.error(f"JSON Parsing Error: {e}")
            status_text.error(f"JSON Parsing Error on Attempt {attempt + 1}")
            logger.error(f"Problematic content: {raw_content}")
            
            # Additional logging to help diagnose the issue
            try:
                import traceback
                logger.error(f"Traceback: {traceback.format_exc()}")
            except Exception:
                pass
        
        except Exception as e:
            logger.error(f"Unexpected Error: {e}")
            status_text.error(f"Unexpected Error on Attempt {attempt + 1}")
            
            # Additional logging
            try:
                import traceback
                logger.error(f"Full Traceback: {traceback.format_exc()}")
            except Exception:
                pass
        
    progress_bar.empty()
    st.error("Failed to generate worksheet after maximum retries. Please try again.")
    return None
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=st.secrets["OPENROUTER_API_KEY"]
)

# Prompt Generation
def prompt(subject, topic):
    return f"""Generate a worksheet JSON for {subject} on {topic}:

REQUIREMENTS:
- Create a JSON object with a 'worksheet' array
- Exactly 9 elements in the array
- First element: Worksheet title
- Next 5 elements: Comprehension/knowledge questions
- Next 2 elements: Multiple-choice questions
- Last element: Cloze passage

EXAMPLE FORMAT:
{{
    "worksheet": [
        "{subject} Worksheet: {topic}",
        "Define the primary concept of {topic}",
        "Explain the significance of a key aspect in {topic}",
        "Analyze the relationship between two core ideas",
        "Describe the main characteristics of the subject",
        "Compare and contrast different perspectives",
        "Multiple choice: Which statement best describes...",
        "Multiple choice: Select the correct explanation for...",
        "Complete the passage about {topic} by filling in the blanks..."
    ]
}}

Be specific, educational, and ensure clear, engaging questions."""

def word(worksheet, subject):
    """
    Generate a Word document in memory without file system storage
    """
    try:
        doc = Document()
        
        # Consistent section margins
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # Styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        
        # Header
        current_date = datetime.now()
        date = current_date.strftime("%d.%m.%y")
        
        # Add header
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = f"{subject}\t\t{date}"
        
        # Main title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(worksheet[0])
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.name = 'Calibri'

        # Add some spacing
        doc.add_paragraph()

        # Questions
        for i, content in enumerate(worksheet[1:], start=1):
            question_paragraph = doc.add_paragraph(f"{i}) {content}")
            question_run = question_paragraph.runs[0]
            question_run.font.bold = True
            question_run.font.size = Pt(14)
            
            # Add answer space
            doc.add_paragraph()  # Blank line for answers
        
        # Save to BytesIO instead of file system
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        return bio
    
    except Exception as e:
        st.error(f"Error generating Word document: {e}")
        return None

def create_worksheet(subject_selection, topic):
    # Generate initial worksheet
    worksheet = response(prompt(subject_selection, topic))
    
    if worksheet:
        # Save worksheet to database
        worksheet_id = worksheet_db.save_worksheet(
            subject=subject_selection, 
            topic=topic, 
            worksheet_content=worksheet
        )
        
        # Create Word document in memory
        doc_download = word(worksheet, subject_selection)
        
        if not doc_download:
            st.error("Failed to generate Word document")
            return None, None
        
        # Columns for download and edit
        col1, col2 = st.columns(2)
        
        with col1:
            # Download button
            st.download_button(
                label="Download Worksheet",
                data=doc_download,
                file_name="Worksheet.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{worksheet_id}"
            )
        
        with col2:
            # Edit button
            if st.button("Edit Worksheet", key=f"edit_{worksheet_id}"):
                edit_worksheet(worksheet_id, subject_selection, worksheet)
        
        return doc_download, worksheet_id
    return None, None

def save_worksheet_changes(worksheet_id, subject, edited_worksheet):
    """
    Save worksheet changes to the database
    """
    try:
        with sqlite3.connect('worksheets.db') as conn:
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE worksheets 
                SET worksheet_content = ? 
                WHERE id = ?
            ''', (json.dumps(edited_worksheet), worksheet_id))
            conn.commit()
        
        # Generate updated Word document in memory
        updated_doc = word(edited_worksheet, subject)
        
        return updated_doc
    except Exception as e:
        st.error(f"Error saving worksheet: {e}")
        return None

def worksheet_library_page():
    st.header("Worksheet Library")
    
    # Filter options
    filter_subject = st.selectbox(
        "Filter by Subject", 
        ["All Subjects"] + [
            "Mathematics", "English", "History", "Geography", 
            "Biology", "Chemistry", "Physics", "Computer Science"
        ]
    )
    
    # Retrieve worksheets
    if filter_subject == "All Subjects":
        worksheets = worksheet_db.list_worksheets(limit=20)
    else:
        worksheets = worksheet_db.list_worksheets(subject=filter_subject, limit=20)
    
    # Display worksheets
    for idx, worksheet in enumerate(worksheets):
        worksheet_id, subject, topic, created_at = worksheet
        
        with st.expander(f"{subject} - {topic} (Created: {created_at})"):
            # Retrieve full worksheet details
            full_worksheet = worksheet_db.get_worksheet_by_id(worksheet_id)
            
            # Display worksheet content
            for i, item in enumerate(full_worksheet['worksheet_content'], 1):
                st.write(f"{i}. {item}")
            
            # Generate Word document in memory
            doc_bytes = word(full_worksheet['worksheet_content'], subject)
            
            if doc_bytes:
                # Download button with unique key
                st.download_button(
                    label="Download Worksheet",
                    data=doc_bytes,
                    file_name=f"{subject}_{topic}_worksheet.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_library_{idx}"
                )
            else:
                st.error(f"Failed to generate Word document for {subject} - {topic}")
    """
    Save worksheet changes to the database with comprehensive error handling
    """
    try:
        # Validate input
        if not worksheet_id or not subject or not edited_worksheet:
            st.error("Invalid input: Missing worksheet details")
            return None

        with sqlite3.connect('worksheets.db') as conn:
            cursor = conn.cursor()
            
            # Check if worksheet exists
            cursor.execute('SELECT * FROM worksheets WHERE id = ?', (worksheet_id,))
            existing_worksheet = cursor.fetchone()
            
            if not existing_worksheet:
                st.error(f"No worksheet found with ID: {worksheet_id}")
                return None

            # Update worksheet content
            cursor.execute('''
                UPDATE worksheets 
                SET worksheet_content = ?, 
                    subject = ?
                WHERE id = ?
            ''', (json.dumps(edited_worksheet), subject, worksheet_id))
            
            # Ensure changes are committed
            conn.commit()
        
        # Generate updated Word document
        doc_download = word(edited_worksheet, subject)
        
        if doc_download:
            st.success("Worksheet updated successfully!")
            return doc_download
        else:
            st.error("Failed to generate updated Word document")
            return None
    
    except sqlite3.Error as e:
        st.error(f"Database error: {e}")
        logger.error(f"Database error in save_worksheet_changes: {e}")
    except Exception as e:
        st.error(f"Unexpected error saving worksheet: {e}")
        logger.error(f"Unexpected error in save_worksheet_changes: {e}")
    
    return None
def edit_worksheet(worksheet_id, subject, worksheet_content):
    """
    Create a comprehensive editable view of the worksheet with versioning
    """
    st.header(f"Edit Worksheet: {subject}")
    
    # Retrieve original worksheet
    original_worksheet = worksheet_content.copy()
    
    # Create a form to contain all editable elements
    with st.form(key=f"edit_worksheet_form_{worksheet_id}"):
        # Editable text areas for each question
        edited_worksheet = []
        for i, content in enumerate(worksheet_content):
            edited_item = st.text_area(
                f"Question {i+1}", 
                value=content, 
                key=f"edit_item_{worksheet_id}_{i}"
            )
            edited_worksheet.append(edited_item)
        
        # Submit button for saving changes
        save_changes = st.form_submit_button("Save Changes")
        
        # Handle saving changes
        if save_changes:
            try:
                # Generate a new unique ID for the edited worksheet
                new_worksheet_id = str(uuid.uuid4())
                
                # Current timestamp
                created_at = datetime.now()
                
                # Save as new worksheet, preserving original
                with sqlite3.connect('worksheets.db') as conn:
                    cursor = conn.cursor()
                    cursor.execute('''
                        INSERT INTO worksheets 
                        (id, subject, topic, worksheet_content, created_at) 
                        VALUES (?, ?, ?, ?, ?)
                    ''', (
                        new_worksheet_id, 
                        subject, 
                        f"Edited Version of {worksheet_id}", 
                        json.dumps(edited_worksheet), 
                        created_at
                    ))
                    conn.commit()
                
                # Generate updated Word document
                updated_doc = word(edited_worksheet, subject)
                
                if updated_doc:
                    st.success(f"Worksheet updated as new version. New Worksheet ID: {new_worksheet_id}")
                    
                    # Download button for updated worksheet
                    st.download_button(
                        label="Download Updated Worksheet",
                        data=updated_doc,
                        file_name=f"{subject}_updated_worksheet.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_updated_{new_worksheet_id}"
                    )
                else:
                    st.error("Failed to generate updated document")
            
            except Exception as e:
                st.error(f"Error saving worksheet: {e}")
                logger.error(f"Worksheet save error: {e}")
def modify_worksheets_table():
    """
    Ensure worksheet_content is stored as text with comprehensive logging
    """
    try:
        with sqlite3.connect('worksheets.db') as conn:
            cursor = conn.cursor()
            
            # More detailed column inspection
            cursor.execute('''
                PRAGMA table_info(worksheets)
            ''')
            columns = cursor.fetchall()
            
            # Detailed logging of existing columns
            logger.info(f"Existing columns: {columns}")
            
            # Check if worksheet_content column exists and is of correct type
            worksheet_content_column = [col for col in columns if col[1] == 'worksheet_content']
            
            if not worksheet_content_column:
                # If column doesn't exist, add it with precise logging
                logger.warning("worksheet_content column missing. Adding column.")
                cursor.execute('''
                    ALTER TABLE worksheets 
                    ADD COLUMN worksheet_content TEXT
                ''')
                logger.info("worksheet_content column added successfully")
            else:
                logger.info("worksheet_content column already exists")
            
            conn.commit()
    except Exception as e:
        logger.error(f"Error modifying worksheets table: {e}")
        st.error(f"Database modification error: {e}")
# Call this method when initializing the database
modify_worksheets_table()


def database_viewer_page():
    st.header("📊 Worksheet Database Viewer")
    
    # Database connection method
    def connect_db():
        return sqlite3.connect('worksheets.db')
    
    # Tabs for different views
    tab1, tab2, tab3 = st.tabs(["📋 All Worksheets", "📊 Database Stats", "🔍 Raw Database Query"])
    
    with tab1:
        st.subheader("All Stored Worksheets")
        
        # Retrieve all worksheets
        with connect_db() as conn:
            df = pd.read_sql_query("""
                SELECT 
                    id, 
                    subject, 
                    topic, 
                    created_at, 
                    length(worksheet_content) as content_length
                FROM worksheets 
                ORDER BY created_at DESC
            """, conn)
        
        # Display as table
        st.dataframe(df, use_container_width=True)
        
        # Detailed view option
        if st.checkbox("Show Detailed Worksheet Contents"):
            worksheet_id = st.selectbox("Select Worksheet ID", df['id'].tolist())
            
            with connect_db() as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM worksheets WHERE id = ?", (worksheet_id,))
                worksheet = cursor.fetchone()
            
            if worksheet:
                st.json({
                    "ID": worksheet[0],
                    "Subject": worksheet[1],
                    "Topic": worksheet[2],
                    "Worksheet Content": json.loads(worksheet[3]),
                    "Created At": worksheet[4],
                    "User ID": worksheet[5]
                })
    
    with tab2:
        st.subheader("Database Statistics")
        
        with connect_db() as conn:
            # Total worksheets
            total_worksheets = pd.read_sql_query("SELECT COUNT(*) as count FROM worksheets", conn).iloc[0,0]
            
            # Worksheets by subject
            subject_counts = pd.read_sql_query("""
                SELECT subject, COUNT(*) as count 
                FROM worksheets 
                GROUP BY subject 
                ORDER BY count DESC
            """, conn)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("Total Worksheets", total_worksheets)
        
        with col2:
            st.write("Worksheets by Subject:")
            st.dataframe(subject_counts, use_container_width=True)
    
    with tab3:
        st.subheader("Raw SQL Query")
        
        # SQL query input
        query = st.text_area("Enter SQL Query", 
                              "SELECT * FROM worksheets LIMIT 10", 
                              height=100)
        
        if st.button("Execute Query"):
            try:
                with connect_db() as conn:
                    df = pd.read_sql_query(query, conn)
                st.dataframe(df, use_container_width=True)
            except Exception as e:
                st.error(f"Error executing query: {e}")

# Worksheet Library Page
# In the worksheet_library_page() function, modify the download button:
def worksheet_library_page():
    st.header("Worksheet Library")
    
    # Filter options
    filter_subject = st.selectbox(
        "Filter by Subject", 
        ["All Subjects"] + [
            "Mathematics", "English", "History", "Geography", 
            "Biology", "Chemistry", "Physics", "Computer Science"
        ]
    )
    
    # Retrieve worksheets
    if filter_subject == "All Subjects":
        worksheets = worksheet_db.list_worksheets(limit=20)
    else:
        worksheets = worksheet_db.list_worksheets(subject=filter_subject, limit=20)
    
    # Display worksheets
    for idx, worksheet in enumerate(worksheets):
        worksheet_id, subject, topic, created_at = worksheet
        
        with st.expander(f"{subject} - {topic} (Created: {created_at})"):
            # Retrieve full worksheet details
            full_worksheet = worksheet_db.get_worksheet_by_id(worksheet_id)
            
            # Display worksheet content
            for i, item in enumerate(full_worksheet['worksheet_content'], 1):
                st.write(f"{i}. {item}")
            
            # Columns for actions
            col1, col2 = st.columns(2)
            
            with col1:
                # Generate Word document in memory
                doc_bytes = word(full_worksheet['worksheet_content'], subject)
                
                if doc_bytes:
                    # Download button with unique key
                    st.download_button(
                        label="Download Worksheet",
                        data=doc_bytes,
                        file_name=f"{subject}_{topic}_worksheet.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_library_{idx}"
                    )
            
            with col2:
                # Edit button
                if st.button(f"Edit Worksheet", key=f"edit_library_{worksheet_id}"):
                    edit_worksheet(
                        worksheet_id, 
                        subject, 
                        full_worksheet['worksheet_content']
                    )
# In the main() function, modify the download button:
def main():
    # ... existing code ...
    if topic and create_button:
        doc_download, worksheet_id = create_worksheet(subject_selection, topic)
        
        if doc_download:
            bio = io.BytesIO()
            doc = Document(doc_download)
            doc.save(bio)

            st.success(f"Worksheet created successfully! ID: {worksheet_id}")

            st.download_button(
                label="Click here to download",
                data=bio.getvalue(),
                file_name="Worksheet.docx",
                mime="docx",
                key="download_main_worksheet"  # Add a unique key
            )
        else:
            st.error("Failed to create worksheet. Please try again.")
# Main Application
def main():
    st.set_page_config(
        page_title="CASE - Worksheet Generator",
        page_icon="🏫",
    )

    page = st.sidebar.radio("Navigate", [
        "Create Worksheet", 
        "Worksheet Library", 
        "📊 Database Viewer"
    ])
    
    if page == "Create Worksheet":
        st.title("CASE - Worksheet Generator")
        
        subject_selection = st.selectbox(
            "Subject",
            ["Select a Subject"] + ["Enter Your Own Subject"] + [
                "Mathematics 🔢", "English 🇬🇧", "History 📜", "Geography 🌍", 
                "Biology 🌿", "Chemistry 🧪", "Physics ⚙️", 
                "Computer Science 💻", "Music 🎵", "Art 🎨", 
                "Sports 🏃‍♂️", "Ethics 🤔", "Religion ⛪", 
                "Politics 🗳️", "Economics 💹", "Philosophy 🤯", 
                "Social Studies 👥", "Psychology 🧠", "Sociology 👩‍👩‍👧‍👦", 
                "Foreign Language 🗣️", "Latin 🏛️", 
                "Spanish 🇪🇸", "French 🇫🇷", "Italian 🇮🇹", 
                "Russian 🇷🇺", "Chinese 🇨🇳", "Japanese 🇯🇵", 
                "Korean 🇰🇷", "Arabic 🇸🇦", "Media Studies 📱"
            ]
        )

        if subject_selection != "Select a Subject":
            if subject_selection == "Enter Your Own Subject":
                subject_selection = st.text_input("Enter Your Own Subject")[:-1]
                st.empty()

            topic = st.text_input("Topic:")

            create_button = st.button("Create Worksheet")

            if topic and create_button:
                doc_download, worksheet_id = create_worksheet(subject_selection, topic)
                
                if doc_download:
                    bio = io.BytesIO()
                    doc = Document(doc_download)
                    doc.save(bio)

                    st.success(f"Worksheet created successfully! ID: {worksheet_id}")

                    st.download_button(
                        label="Click here to download",
                        data=bio.getvalue(),
                        file_name="Worksheet.docx",
                        mime="docx",
                        key="download_button"
                    )
                else:
                    st.error("Failed to create worksheet. Please try again.")
    
    elif page == "Worksheet Library":
        worksheet_library_page()
    
    elif page == "📊 Database Viewer":
        database_viewer_page()

# Run the application
if __name__ == "__main__":
    main()